#!/usr/bin/env python3
"""
Textract Service for Document Text Extraction

This service provides comprehensive functionality for extracting text from various
document formats using local processing libraries.

Features:
- PDF text extraction using PyPDF2
- DOCX text extraction using python-docx
- TXT file processing
- Text cleaning and preprocessing
- Support for multiple file formats
"""

import os
import logging
import mimetypes
from typing import Dict, Any, Optional, List
from pathlib import Path
import PyPDF2
from docx import Document
import dotenv

# Load environment variables
dotenv.load_dotenv()

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class TextractService:
    """
    Service for extracting text from various document formats.
    
    This service handles text extraction from PDFs, DOCX files, TXT files.
    """
    
    def extract_text_from_file(self, file_path: str) -> Dict[str, Any]:
        """
        Extract text from a file using the appropriate method.
        
        Args:
            file_path: Path to the file to extract text from
            
        Returns:
            Dict containing extracted text and metadata
        """
        try:
            if not os.path.exists(file_path):
                return {'success': False, 'error': 'File not found'}
            
            # Get file info
            file_info = self._get_file_info(file_path)
            if not file_info['success']:
                return file_info
            
            file_type = file_info['file_type']
            file_size = file_info['file_size']
            
            # Extract text based on file type
            if file_type == 'pdf':
                result = self._extract_text_from_pdf_local(file_path)
            elif file_type == 'docx':
                result = self._extract_text_from_docx(file_path)
            elif file_type == 'txt':
                result = self._extract_text_from_txt(file_path)
            else:
                return {'success': False, 'error': f'Unsupported file type: {file_type}'}
            
            if result['success']:
                # Add file metadata
                result.update({
                    'file_path': file_path,
                    'file_type': file_type,
                    'file_size': file_size,
                    'extraction_method': result.get('extraction_method', 'local')
                })
            
            return result
            
        except Exception as e:
            logger.error(f"Error extracting text from file {file_path}: {str(e)}")
            return {'success': False, 'error': str(e)}
    
    def extract_text_from_bytes(self, file_bytes: bytes, filename: str) -> Dict[str, Any]:
        """
        Extract text from file bytes.
        
        Args:
            file_bytes: File content as bytes
            filename: Original filename (used to determine file type)
            
        Returns:
            Dict containing extracted text and metadata
        """
        try:
            # Determine file type from filename
            file_type = self._get_file_type_from_filename(filename)
            if not file_type:
                return {'success': False, 'error': f'Unsupported file type for: {filename}'}
            
            # Extract text based on file type
            if file_type == 'pdf':
                result = self._extract_text_from_pdf_bytes_local(file_bytes)
            elif file_type == 'docx':
                result = self._extract_text_from_docx_bytes(file_bytes)
            elif file_type == 'txt':
                result = self._extract_text_from_txt_bytes(file_bytes)
            else:
                return {'success': False, 'error': f'Unsupported file type: {file_type}'}
            
            if result['success']:
                # Add file metadata
                result.update({
                    'filename': filename,
                    'file_type': file_type,
                    'file_size': len(file_bytes),
                    'extraction_method': result.get('extraction_method', 'local')
                })
            
            return result
            
        except Exception as e:
            logger.error(f"Error extracting text from bytes: {str(e)}")
            return {'success': False, 'error': str(e)}
    
    def _get_file_info(self, file_path: str) -> Dict[str, Any]:
        """Get file information."""
        try:
            stat = os.stat(file_path)
            file_type = self._get_file_type_from_filename(file_path)
            
            return {
                'success': True,
                'file_type': file_type,
                'file_size': stat.st_size,
                'file_path': file_path
            }
        except Exception as e:
            return {'success': False, 'error': str(e)}
    
    def _get_file_type_from_filename(self, filename: str) -> Optional[str]:
        """Determine file type from filename."""
        ext = Path(filename).suffix.lower()
        
        type_mapping = {
            '.pdf': 'pdf',
            '.docx': 'docx',
            '.doc': 'docx',  # Treat .doc as .docx
            '.txt': 'txt',
            '.rtf': 'txt'   # Treat .rtf as .txt
        }
        
        return type_mapping.get(ext)
    
    def _extract_text_from_pdf_local(self, file_path: str) -> Dict[str, Any]:
        """Extract text from PDF using PyPDF2."""
        try:
            text = ""
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    page_text = page.extract_text()
                    text += page_text + "\n"
            
            # Clean up text
            cleaned_text = self._clean_text(text)
            
            return {
                'success': True,
                'text': cleaned_text,
                'extraction_method': 'local_pypdf2',
                'page_count': len(pdf_reader.pages)
            }
            
        except Exception as e:
            logger.error(f"Error extracting text from PDF: {str(e)}")
            return {'success': False, 'error': str(e)}
    
    def _extract_text_from_pdf_bytes_local(self, file_bytes: bytes) -> Dict[str, Any]:
        """Extract text from PDF bytes using PyPDF2."""
        try:
            import io
            
            text = ""
            pdf_file = io.BytesIO(file_bytes)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                page_text = page.extract_text()
                text += page_text + "\n"
            
            # Clean up text
            cleaned_text = self._clean_text(text)
            
            return {
                'success': True,
                'text': cleaned_text,
                'extraction_method': 'local_pypdf2',
                'page_count': len(pdf_reader.pages)
            }
            
        except Exception as e:
            logger.error(f"Error extracting text from PDF bytes: {str(e)}")
            return {'success': False, 'error': str(e)}
    
    
    def _extract_text_from_docx(self, file_path: str) -> Dict[str, Any]:
        """Extract text from DOCX file."""
        try:
            doc = Document(file_path)
            text = ""
            
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Clean up text
            cleaned_text = self._clean_text(text)
            
            return {
                'success': True,
                'text': cleaned_text,
                'extraction_method': 'local_docx'
            }
            
        except Exception as e:
            logger.error(f"Error extracting text from DOCX: {str(e)}")
            return {'success': False, 'error': str(e)}
    
    def _extract_text_from_docx_bytes(self, file_bytes: bytes) -> Dict[str, Any]:
        """Extract text from DOCX bytes."""
        try:
            import io
            
            doc = Document(io.BytesIO(file_bytes))
            text = ""
            
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Clean up text
            cleaned_text = self._clean_text(text)
            
            return {
                'success': True,
                'text': cleaned_text,
                'extraction_method': 'local_docx'
            }
            
        except Exception as e:
            logger.error(f"Error extracting text from DOCX bytes: {str(e)}")
            return {'success': False, 'error': str(e)}
    
    def _extract_text_from_txt(self, file_path: str) -> Dict[str, Any]:
        """Extract text from TXT file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
            
            # Clean up text
            cleaned_text = self._clean_text(text)
            
            return {
                'success': True,
                'text': cleaned_text,
                'extraction_method': 'local_txt'
            }
            
        except UnicodeDecodeError:
            # Try with different encoding
            try:
                with open(file_path, 'r', encoding='latin-1') as file:
                    text = file.read()
                
                cleaned_text = self._clean_text(text)
                
                return {
                    'success': True,
                    'text': cleaned_text,
                    'extraction_method': 'local_txt_latin1'
                }
            except Exception as e:
                logger.error(f"Error extracting text from TXT with latin-1: {str(e)}")
                return {'success': False, 'error': str(e)}
        except Exception as e:
            logger.error(f"Error extracting text from TXT: {str(e)}")
            return {'success': False, 'error': str(e)}
    
    def _extract_text_from_txt_bytes(self, file_bytes: bytes) -> Dict[str, Any]:
        """Extract text from TXT bytes."""
        try:
            text = file_bytes.decode('utf-8')
            
            # Clean up text
            cleaned_text = self._clean_text(text)
            
            return {
                'success': True,
                'text': cleaned_text,
                'extraction_method': 'local_txt'
            }
            
        except UnicodeDecodeError:
            # Try with different encoding
            try:
                text = file_bytes.decode('latin-1')
                
                cleaned_text = self._clean_text(text)
                
                return {
                    'success': True,
                    'text': cleaned_text,
                    'extraction_method': 'local_txt_latin1'
                }
            except Exception as e:
                logger.error(f"Error extracting text from TXT bytes with latin-1: {str(e)}")
                return {'success': False, 'error': str(e)}
        except Exception as e:
            logger.error(f"Error extracting text from TXT bytes: {str(e)}")
            return {'success': False, 'error': str(e)}
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize extracted text."""
        if not text:
            return ""
        
        # Remove excessive whitespace
        text = ' '.join(text.split())
        
        # Remove special characters that might cause issues
        text = text.replace('\x00', '')  # Remove null bytes
        
        # Normalize line breaks
        text = text.replace('\r\n', '\n').replace('\r', '\n')
        
        return text.strip()

    
    def get_supported_formats(self) -> List[str]:
        """Get list of supported file formats."""
        return ['pdf', 'docx', 'doc', 'txt', 'rtf']
    
    def is_format_supported(self, filename: str) -> bool:
        """Check if file format is supported."""
        file_type = self._get_file_type_from_filename(filename)
        return file_type is not None
